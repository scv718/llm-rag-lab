import io
import csv
import traceback

import PyPDF2
from fastapi import HTTPException

from app.core.config import DOCS_DIR, REPO_DIR, UPLOAD_DIR, vs
from app.core.state import latest_target
from app.db.target_repo import upsert_project_target
from app.services.chunking import TextChunk, chunk_text_by_chars
from app.services.embeddings import embed_documents
from app.services.zip_ingest import (
    CodeChunk,
    chunk_code_by_lines,
    decode_text_best_effort,
    _guess_lang,
    iter_source_files,
    read_text_best_effort,
    safe_extract_zip,
)
from app.utils.file_utils import detect_file_kind
from app.utils.hash_utils import sha256_bytes


TEXT_DOCUMENT_KINDS = {"txt", "md", "text"}
CODE_FILE_SOURCE_TYPES = {"java", "python", "javascript", "typescript", "xml", "yaml", "sql"}


def ingest_upload(raw: bytes, filename: str, project_id: int | None = None) -> dict:
    kind = detect_file_kind(raw, filename)
    if kind == "unknown":
        raise HTTPException(
            status_code=400,
            detail="지원하지 않는 파일 형식입니다. (PDF/ZIP/TXT/MD/DOCX/XLSX/CSV/단일 코드파일 지원)",
        )

    try:
        if kind == "pdf":
            return ingest_pdf(raw, filename, project_id=project_id)
        if kind == "zip":
            return ingest_zip(raw, filename, project_id=project_id)
        if kind in TEXT_DOCUMENT_KINDS:
            return ingest_text_document(raw, filename, source_type=kind, project_id=project_id)
        if kind == "csv":
            return ingest_csv(raw, filename, project_id=project_id)
        if kind == "docx":
            return ingest_docx(raw, filename, project_id=project_id)
        if kind == "xlsx":
            return ingest_xlsx(raw, filename, project_id=project_id)
        if kind == "code":
            return ingest_code_file(raw, filename, project_id=project_id)
    except HTTPException:
        raise
    except Exception as exc:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    raise HTTPException(status_code=400, detail="알 수 없는 업로드 유형입니다.")


def ingest_pdf(raw: bytes, filename: str, project_id: int | None = None) -> dict:
    doc_id = f"doc-{sha256_bytes(raw)[:16]}"
    pdf_path = DOCS_DIR / f"{doc_id}.pdf"
    pdf_path.write_bytes(raw)

    pdf_reader = PyPDF2.PdfReader(io.BytesIO(raw))
    all_chunks: list[TextChunk] = []

    for i, page in enumerate(pdf_reader.pages):
        page_no = i + 1
        page_text = page.extract_text() or ""

        page_chunks = chunk_text_by_chars(
            doc_id=doc_id,
            page=page_no,
            text=page_text,
            chunk_size=900,
            overlap=150,
        )
        all_chunks.extend(page_chunks)

    return finalize_doc_ingest(
        doc_id=doc_id,
        filename=filename,
        source_type="pdf",
        project_id=project_id,
        chunks=all_chunks,
        stored_path=str(pdf_path),
        extra={"pages": len(pdf_reader.pages)},
    )


def ingest_zip(raw: bytes, filename: str, project_id: int | None = None) -> dict:
    repo_id = sha256_bytes(raw)[:16]
    zip_path = UPLOAD_DIR / f"{repo_id}.zip"
    zip_path.write_bytes(raw)

    extract_dir = REPO_DIR / repo_id
    if extract_dir.exists():
        import shutil

        shutil.rmtree(extract_dir, ignore_errors=True)

    extracted_files, total_bytes = safe_extract_zip(
        zip_path=zip_path,
        extract_dir=extract_dir,
        max_files=8000,
        max_total_uncompressed_bytes=300 * 1024 * 1024,
    )

    all_chunks: list[CodeChunk] = []
    for src in iter_source_files(extract_dir):
        rel_path = str(src.relative_to(extract_dir)).replace("\\", "/")
        text = read_text_best_effort(src)
        all_chunks.extend(
            chunk_code_by_lines(
                repo_id=repo_id,
                rel_path=rel_path,
                text=text,
                lines_per_chunk=250,
                overlap=40,
            )
        )

    finalize_code_ingest(
        repo_id=repo_id,
        filename=filename,
        source_type="zip",
        project_id=project_id,
        chunks=all_chunks,
    )

    return {
        "status": "success",
        "kind": "code",
        "repo_id": repo_id,
        "filename": filename,
        "project_id": project_id,
        "extracted_files": extracted_files,
        "uncompressed_bytes": total_bytes,
        "chunks": len(all_chunks),
        "indexed": len(all_chunks),
        "zip_path": str(zip_path),
        "extract_path": str(extract_dir),
    }


def ingest_text_document(
    raw: bytes,
    filename: str,
    source_type: str,
    project_id: int | None = None,
) -> dict:
    doc_id = f"doc-{sha256_bytes(raw)[:16]}"
    suffix = source_type if source_type != "text" else "txt"
    stored_path = DOCS_DIR / f"{doc_id}.{suffix}"
    stored_path.write_bytes(raw)

    text = decode_text_best_effort(raw)
    chunks = chunk_text_by_chars(doc_id=doc_id, page=1, text=text, chunk_size=900, overlap=150)

    return finalize_doc_ingest(
        doc_id=doc_id,
        filename=filename,
        source_type=source_type,
        project_id=project_id,
        chunks=chunks,
        stored_path=str(stored_path),
        extra={"pages": 1},
    )


def ingest_csv(raw: bytes, filename: str, project_id: int | None = None) -> dict:
    doc_id = f"doc-{sha256_bytes(raw)[:16]}"
    stored_path = DOCS_DIR / f"{doc_id}.csv"
    stored_path.write_bytes(raw)

    text = decode_text_best_effort(raw)
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        chunks = []
    else:
        header = rows[0]
        lines = []
        for idx, row in enumerate(rows[1:], start=2):
            pairs = []
            for col_idx, value in enumerate(row):
                key = header[col_idx].strip() if col_idx < len(header) and header[col_idx].strip() else f"col_{col_idx + 1}"
                pairs.append(f"{key}: {value}")
            if pairs:
                lines.append(f"row {idx} | " + " | ".join(pairs))

        body = "\n".join(lines) if lines else "\n".join([",".join(row) for row in rows])
        chunks = chunk_text_by_chars(doc_id=doc_id, page=1, text=body, chunk_size=1200, overlap=120)

    return finalize_doc_ingest(
        doc_id=doc_id,
        filename=filename,
        source_type="csv",
        project_id=project_id,
        chunks=chunks,
        stored_path=str(stored_path),
        extra={"pages": 1, "rows": max(0, len(rows) - 1) if rows else 0},
    )


def ingest_docx(raw: bytes, filename: str, project_id: int | None = None) -> dict:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="DOCX 처리를 위해 python-docx 설치가 필요합니다.") from exc

    doc_id = f"doc-{sha256_bytes(raw)[:16]}"
    stored_path = DOCS_DIR / f"{doc_id}.docx"
    stored_path.write_bytes(raw)

    document = Document(io.BytesIO(raw))
    blocks = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row_idx, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(f"table row {row_idx} | " + " | ".join(cells))

    text = "\n\n".join(blocks)
    chunks = chunk_text_by_chars(doc_id=doc_id, page=1, text=text, chunk_size=1000, overlap=150)

    return finalize_doc_ingest(
        doc_id=doc_id,
        filename=filename,
        source_type="docx",
        project_id=project_id,
        chunks=chunks,
        stored_path=str(stored_path),
        extra={"pages": 1},
    )


def ingest_xlsx(raw: bytes, filename: str, project_id: int | None = None) -> dict:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="XLSX 처리를 위해 openpyxl 설치가 필요합니다.") from exc

    doc_id = f"doc-{sha256_bytes(raw)[:16]}"
    stored_path = DOCS_DIR / f"{doc_id}.xlsx"
    stored_path.write_bytes(raw)

    workbook = load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
    sheet_count = len(workbook.sheetnames)
    all_chunks: list[TextChunk] = []

    for sheet_index, sheet_name in enumerate(workbook.sheetnames, start=1):
        sheet = workbook[sheet_name]
        rows = []
        header = None

        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = ["" if value is None else str(value).strip() for value in row]
            if not any(values):
                continue

            if header is None:
                header = values
                continue

            pairs = []
            for col_idx, value in enumerate(values):
                if not value:
                    continue
                key = header[col_idx].strip() if col_idx < len(header) and header[col_idx].strip() else f"col_{col_idx + 1}"
                pairs.append(f"{key}: {value}")

            if pairs:
                rows.append(f"sheet {sheet_name} row {row_idx} | " + " | ".join(pairs))

        if not rows:
            continue

        sheet_text = "\n".join(rows)
        all_chunks.extend(
            chunk_text_by_chars(
                doc_id=doc_id,
                page=sheet_index,
                text=sheet_text,
                chunk_size=1200,
                overlap=120,
            )
        )

    workbook.close()

    return finalize_doc_ingest(
        doc_id=doc_id,
        filename=filename,
        source_type="xlsx",
        project_id=project_id,
        chunks=all_chunks,
        stored_path=str(stored_path),
        extra={"pages": sheet_count},
    )


def ingest_code_file(raw: bytes, filename: str, project_id: int | None = None) -> dict:
    repo_id = sha256_bytes(raw)[:16]
    extract_dir = REPO_DIR / repo_id
    extract_dir.mkdir(parents=True, exist_ok=True)

    stored_path = extract_dir / filename
    stored_path.parent.mkdir(parents=True, exist_ok=True)
    stored_path.write_bytes(raw)

    text = decode_text_best_effort(raw)
    chunks = chunk_code_by_lines(
        repo_id=repo_id,
        rel_path=filename,
        text=text,
        lines_per_chunk=250,
        overlap=40,
    )

    source_type = _guess_lang(filename)
    if source_type not in CODE_FILE_SOURCE_TYPES:
        source_type = "text"

    finalize_code_ingest(
        repo_id=repo_id,
        filename=filename,
        source_type=source_type,
        project_id=project_id,
        chunks=chunks,
    )

    return {
        "status": "success",
        "kind": "code",
        "repo_id": repo_id,
        "filename": filename,
        "project_id": project_id,
        "extracted_files": 1,
        "uncompressed_bytes": len(raw),
        "chunks": len(chunks),
        "indexed": len(chunks),
        "extract_path": str(extract_dir),
        "stored_path": str(stored_path),
    }


def finalize_doc_ingest(
    *,
    doc_id: str,
    filename: str,
    source_type: str,
    project_id: int | None,
    chunks: list[TextChunk],
    stored_path: str,
    extra: dict | None = None,
) -> dict:
    chunk_texts = [c.text for c in chunks]
    if chunk_texts:
        vectors = embed_documents(chunk_texts)
        ids = [c.chunk_id for c in chunks]
        metadatas = [
            {
                "kind": "doc",
                "source_type": source_type,
                "doc_id": doc_id,
                "filename": filename,
                "page": c.page,
                "chunk_index": c.chunk_index,
                "chunk_id": c.chunk_id,
            }
            for c in chunks
        ]
        vs.upsert(ids=ids, embeddings=vectors, metadatas=metadatas, documents=chunk_texts)

    latest_target["kind"] = "doc"
    latest_target["id"] = doc_id
    latest_target["filename"] = filename

    if project_id is not None:
        upsert_project_target(project_id, "doc", doc_id, filename)

    response = {
        "status": "success",
        "kind": "doc",
        "doc_id": doc_id,
        "filename": filename,
        "project_id": project_id,
        "source_type": source_type,
        "chunks": len(chunks),
        "indexed": len(chunks),
        "stored_path": stored_path,
    }
    if extra:
        response.update(extra)
    return response


def finalize_code_ingest(
    *,
    repo_id: str,
    filename: str,
    source_type: str,
    project_id: int | None,
    chunks: list[CodeChunk],
) -> None:
    texts = [c.text for c in chunks]
    if texts:
        vectors = embed_documents(texts)

        ids = [c.chunk_id for c in chunks]
        metadatas = [
            {
                "kind": "code",
                "source_type": source_type,
                "repo_id": repo_id,
                "filename": filename,
                "path": c.path,
                "lang": c.lang,
                "start_line": c.start_line,
                "end_line": c.end_line,
                "chunk_id": c.chunk_id,
            }
            for c in chunks
        ]
        vs.upsert(ids=ids, embeddings=vectors, metadatas=metadatas, documents=texts)

    latest_target["kind"] = "code"
    latest_target["id"] = repo_id
    latest_target["filename"] = filename

    if project_id is not None:
        upsert_project_target(project_id, "code", repo_id, filename)
